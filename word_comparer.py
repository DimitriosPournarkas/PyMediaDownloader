import sys
import os
from docx import Document
import re

def compare_word_files(file1, file2):
    try:
        
        # Read Word documents
        doc1 = Document(file1)
        doc2 = Document(file2)
        
        # Extract text from both documents
        text1 = extract_text(doc1)
        text2 = extract_text(doc2)
        
        # Calculate similarity
        similarity = calculate_text_similarity(text1, text2)
        
        return similarity > 0.6  # 60% content similarity threshold
        
    except Exception as e:
        
        return False

def extract_text(doc):
    """Extract text from Word document"""
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
    
    return text.strip()

def calculate_text_similarity(text1, text2):
    """Calculate similarity between two texts"""
    if not text1 or not text2:
        return 0.0
    
    # Simple word-based similarity
    words1 = set(re.findall(r'\w+', text1.lower()))
    words2 = set(re.findall(r'\w+', text2.lower()))
    
    if not words1 or not words2:
        return 0.0
    
    common_words = words1.intersection(words2)
    similarity = len(common_words) / max(len(words1), len(words2))
    
    return similarity

if __name__ == "__main__":
    if len(sys.argv) != 3:
        
        sys.exit(1)
    
    file1, file2 = sys.argv[1], sys.argv[2]
    similar = compare_word_files(file1, file2)
    sys.exit(0 if similar else 1)